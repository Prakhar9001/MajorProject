import os
import zlib
import base64
import requests
import re
from docx import Document
from docx.shared import Inches

md_path = r"C:\Users\User\.gemini\antigravity\brain\7577cbf2-26cb-4a4e-b20f-d0ddf89ab7de\MediBot_Architecture_Document.md"
docx_path = r"C:\Users\User\Downloads\MediBot_Architecture_Document.docx"

diagrams = {
    "3.5 System Interaction Summary": "graph TD\nUser([User]) --> |Text Query / Image| UI[Chainlit UI]\nUI --> |Uploads| Vision[Vision Pipeline]\nUI --> |Queries| RAG[RAG Pipeline]\nVision --> CNN[PyTorch CNN]\nCNN --> |Fracture/Disease| UI\nRAG --> Embed[HF Embeddings]\nEmbed --> FAISS[(FAISS Vector DB)]\nFAISS --> |Context Chunks| Prompt[Prompt Template]\nPrompt --> LLM[Llama-2-7B]\nLLM --> |Generated Answer| UI",
    "6.3 Component Dependency Map": "graph TD\nsubgraph Presentation\nUI[Chainlit App]\nend\nsubgraph NLP Layer\nQA[RetrievalQA Chain]\nLLM[CTransformers LLM]\nend\nsubgraph Retrieval Layer\nFAISS[FAISS Index]\nEmbed[MiniLM Embeddings]\nend\nsubgraph Vision Layer\nCNN[PyTorch CNN]\nend\nUI --> QA\nUI --> CNN\nQA --> LLM\nQA --> FAISS\nFAISS --> Embed",
    "9.1 Document Ingestion Sequence": "sequenceDiagram\nparticipant Admin\nparticipant PyPDF as PyPDFLoader\nparticipant Splitter as TextSplitter\nparticipant Embed as HF Embeddings\nparticipant FAISS\nAdmin->>PyPDF: ingest.py\nPyPDF-->>Splitter: Raw Text Document\nSplitter-->>Embed: 600-char Text Chunks\nEmbed-->>FAISS: Vector Embeddings\nFAISS-->>Admin: Saved db_faiss locally",
    "9.2 Text QA Retrieval Sequence": "sequenceDiagram\nparticipant User\nparticipant Chainlit\nparticipant FAISS\nparticipant Llama2\nUser->>Chainlit: Asks Medical Question\nChainlit->>FAISS: Search(Embedded Query)\nFAISS-->>Chainlit: Top 3 Context Chunks\nChainlit->>Llama2: Prompt(Context + Query)\nLlama2-->>Chainlit: Streamed Response\nChainlit-->>User: Display Text Output",
    "9.3 Lung Image Prediction Sequence": "sequenceDiagram\nparticipant User\nparticipant Chainlit\nparticipant CNN Model\nUser->>Chainlit: Uploads X-Ray\nChainlit->>CNN Model: predict_lung_disease(bytes)\nCNN Model->>CNN Model: Resize 224x224 & Normalize\nCNN Model->>CNN Model: PyTorch Forward Pass\nCNN Model->>CNN Model: Softmax Confidence\nCNN Model-->>Chainlit: \"Disease Detected (95%)\"\nChainlit-->>User: Display Prediction Message"
}

images = {}
for name, code in diagrams.items():
    encoded = base64.urlsafe_b64encode(zlib.compress(code.encode('utf-8'), 9)).decode('ascii')
    url = f"https://kroki.io/mermaid/png/{encoded}"
    resp = requests.get(url)
    img_path = f"{name.replace(' ', '_').replace('.', '_')}.png"
    if resp.status_code == 200:
        with open(img_path, 'wb') as f:
            f.write(resp.content)
        images[name] = img_path
    else:
        print(f"Failed to fetch {name}")

doc = Document()
doc.add_heading('Architecture Document: MediBot', 0)

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    # Skip preamble lines that we handle manually
    if line == 'Architecture Document':
        continue
    if line == 'MediBot Medical QA and Disease Detection System':
        continue
    if line == '[Llama-2 RAG Pipeline with PyTorch CNN Evaluation,':
        continue
    if line == 'Local Vector Storage & Interactive Web Interface]':
        continue
        
    m1 = re.match(r'^([0-9]+)\.\s+(.*)', line)
    m2 = re.match(r'^([0-9]+)\.([0-9]+)\s+(.*)', line)
    m3 = re.match(r'^([0-9]+)\.([0-9]+)\.([0-9]+)\s+(.*)', line)

    if line.startswith('Table of Contents'):
        doc.add_heading(line, level=1)
    elif m3:
        doc.add_heading(line, level=3)
    elif m2:
        doc.add_heading(line, level=2)
    elif m1:
        doc.add_heading(line, level=1)
    else:
        doc.add_paragraph(line)
        
    # Insert diagram if we just printed the heading for it
    for key, img_path in images.items():
        if line.startswith(key):
            try:
                # Add a small spacing
                doc.add_paragraph()
                doc.add_picture(img_path, width=Inches(6.0))
                doc.add_paragraph()
            except Exception as e:
                print(f"Failed to add picture {img_path}: {e}")

doc.save(docx_path)
print("SUCCESS")
